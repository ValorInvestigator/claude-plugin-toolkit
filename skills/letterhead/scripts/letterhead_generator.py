"""
Valor Investigations Letterhead Generator
Creates .docx from the Valor letterhead template.

Usage:
    from letterhead_generator import create_letterhead_doc
    create_letterhead_doc(output_path, date_str, recipient_lines, subject, salutation, body_paragraphs, closing_title)
"""

import zipfile
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_letterhead_doc(output_path, date_str, recipient_lines, subject, salutation, body_paragraphs, closing_title, attachments=None, footer_title=None):
    """
    Create a .docx from the Valor letterhead template.

    Args:
        output_path: Full path for output .docx
        date_str: Date string (e.g., "February 18, 2026")
        recipient_lines: List of strings for recipient block
        subject: Subject/RE line text
        salutation: Opening line (e.g., "Western Skies Wellness,")
        body_paragraphs: List of dicts with keys:
            - 'text': paragraph text
            - 'bold': bool (optional, default False)
            - 'numbered': bool (optional, for numbered list items)
            - 'bullet': bool (optional, for bullet points)
            - 'indent': bool (optional, for indented text)
        closing_title: String for title under "Levi Bakke" (e.g., "Legal Health Representative for Fred Schoellig")
        attachments: Optional list of attachment description strings
        footer_title: Optional title for footer (defaults to subject)
    """
    # Try multiple known template locations
    template_candidates = [
        r'C:\Users\Big Levi\OneDrive\Documents\Header Template.dotx',
        r'D:\000000000 Letterhead.docx',
    ]
    template_path = None
    for candidate in template_candidates:
        if os.path.exists(candidate):
            template_path = candidate
            break
    if template_path is None:
        raise FileNotFoundError(
            "Valor letterhead template not found. Checked: "
            + ", ".join(template_candidates)
        )
    tmp_path = output_path + '.tmp.docx'

    if template_path.lower().endswith('.dotx'):
        # Convert .dotx to .docx (fix content type)
        with zipfile.ZipFile(template_path, 'r') as zin:
            with zipfile.ZipFile(tmp_path, 'w') as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == '[Content_Types].xml':
                        data = data.replace(
                            b'wordprocessingml.template.main+xml',
                            b'wordprocessingml.document.main+xml'
                        )
                    zout.writestr(item, data)
    else:
        # .docx template -- just copy it
        import shutil
        shutil.copy2(template_path, tmp_path)

    doc = Document(tmp_path)

    # Clear ALL existing body content (paragraphs AND tables)
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith('}sectPr'):
            continue
        body.remove(child)

    def add_para(text, bold=False, alignment=None, space_after=None, font_size=Pt(12)):
        p = doc.add_paragraph()
        if alignment:
            p.alignment = alignment
        run = p.add_run(text)
        run.font.name = 'Courier New'
        run.font.size = font_size
        if bold:
            run.bold = True
        if space_after is not None:
            from docx.shared import Pt as PtSpacing
            p.paragraph_format.space_after = PtSpacing(space_after)
        return p

    # Date
    add_para(date_str)
    add_para('')  # blank line

    # Recipient block
    for line in recipient_lines:
        add_para(line)
    add_para('')  # blank line

    # Subject line
    add_para(f'RE: {subject}', bold=True)
    add_para('')  # blank line

    # Salutation
    add_para(salutation)
    add_para('')  # blank line

    # Body paragraphs
    for para in body_paragraphs:
        text = para if isinstance(para, str) else para.get('text', '')
        is_bold = False if isinstance(para, str) else para.get('bold', False)
        is_numbered = False if isinstance(para, str) else para.get('numbered', False)
        is_bullet = False if isinstance(para, str) else para.get('bullet', False)

        if is_numbered or is_bullet:
            p = add_para(text)
            p.paragraph_format.left_indent = Inches(0.5)
        elif is_bold:
            add_para(text, bold=True)
        elif text == '':
            add_para('')
        else:
            add_para(text)

    add_para('')  # blank line

    # Closing
    add_para('Sincerely,')
    add_para('')
    add_para('')
    add_para('Levi Bakke', bold=True)
    add_para('Valor Investigations')
    if closing_title:
        add_para(closing_title)

    # Attachments section
    if attachments:
        add_para('')
        add_para('Attached:', bold=True)
        for att in attachments:
            p = add_para(f'\u2022  {att}')
            p.paragraph_format.left_indent = Inches(0.25)

    doc.save(output_path)
    os.remove(tmp_path)
    print(f'Saved: {output_path}')


if __name__ == '__main__':
    import sys
    print("Import this module and call create_letterhead_doc(). See docstring for args.")
